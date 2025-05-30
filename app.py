# AI News Generator + RAG Query System
# Multi-Agent System + Retrieval Augmented Generation

import os
import warnings
import streamlit as st
from datetime import datetime
from pathlib import Path
import json
import markdown
import pdfkit
from io import BytesIO
import faiss
import numpy as np
from typing import List

# Suppress warnings
warnings.filterwarnings('ignore')

# CrewAI imports
from crewai import Agent, Crew, Task
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
from openai import OpenAI

# Configure the page
st.set_page_config(
    page_title="AI News Generator + RAG",
    page_icon="🤖",
    layout="wide"
)

# Initialize output directory with absolute path
current_dir = Path.cwd()
output_dir = current_dir / "output"
output_dir.mkdir(exist_ok=True)

# Debug: Show where files will be saved
st.sidebar.info(f"📁 Files will be saved to:\n{output_dir.absolute()}")

class NewsArticle(BaseModel):
    title: str
    summary: str
    key_points: list
    source: str
    url: str
    relevance_score: int

class BlogPost(BaseModel):
    title: str
    introduction: str
    main_content: str
    conclusion: str
    sources: list

# Initialize tools
@st.cache_resource
def initialize_tools():
    search_tool = SerperDevTool()
    scrape_tool = ScrapeWebsiteTool()
    return search_tool, scrape_tool

# Initialize OpenAI client
@st.cache_resource
def initialize_openai_client(api_key):
    return OpenAI(api_key=api_key)

# RAG Components
class RAGSystem:
    def __init__(self, openai_client):
        self.client = openai_client
        self.chunks = []
        self.index = None
        self.embeddings = []
        self.chat_history = []
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            chunks.append(chunk)
        return chunks
    
    def get_embedding(self, text: str, model: str = "text-embedding-3-small"):
        """Get embedding for text"""
        response = self.client.embeddings.create(
            model=model,
            input=text
        )
        return response.data[0].embedding
    
    def build_index(self, document_text: str):
        """Build FAISS index from document text"""
        try:
            # Chunk the document
            self.chunks = self.chunk_text(document_text)
            st.info(f"📄 Document split into {len(self.chunks)} chunks")
            
            # Create embeddings
            self.embeddings = []
            progress_bar = st.progress(0)
            for i, chunk in enumerate(self.chunks):
                embedding = self.get_embedding(chunk)
                self.embeddings.append(embedding)
                progress_bar.progress((i + 1) / len(self.chunks))
            
            # Build FAISS index
            if self.embeddings:
                dimension = len(self.embeddings[0])
                self.index = faiss.IndexFlatL2(dimension)
                self.index.add(np.array(self.embeddings).astype("float32"))
                st.success(f"✅ RAG index built with {len(self.chunks)} chunks")
                return True
            return False
        except Exception as e:
            st.error(f"Error building RAG index: {str(e)}")
            return False
    
    def retrieve_context(self, query: str, k: int = 3) -> str:
        """Retrieve relevant context for query"""
        if self.index is None:
            return ""
        
        query_embedding = self.get_embedding(query)
        D, I = self.index.search(np.array([query_embedding]).astype("float32"), k)
        return "\n---\n".join([self.chunks[i] for i in I[0] if i < len(self.chunks)])
    
    def generate_response(self, query: str) -> str:
        """Generate response using RAG"""
        try:
            context = self.retrieve_context(query)
            
            messages = [
                {"role": "system", "content": "You are a helpful AI assistant that answers questions about AI news using the provided context. Be accurate, concise, and cite specific information from the context when possible."}
            ]
            
            # Add chat history
            messages.extend(self.chat_history[-6:])  # Keep last 6 messages for context
            
            messages.append({
                "role": "user",
                "content": f"Context from AI News Document:\n{context}\n\nQuestion: {query}"
            })
            
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=messages,
                temperature=0.3,
                max_tokens=500
            )
            
            answer = response.choices[0].message.content
            
            # Update chat history
            self.chat_history.append({"role": "user", "content": query})
            self.chat_history.append({"role": "assistant", "content": answer})
            
            return answer
        except Exception as e:
            return f"Error generating response: {str(e)}"

def create_agents(search_tool, scrape_tool):
    """Create the four agents for news scraping and blog writing"""
    
    # Agent 1: New York Times Scraper
    nyt_agent = Agent(
        role="New York Times AI News Specialist",
        goal="Find and analyze the latest AI news from The New York Times",
        tools=[search_tool, scrape_tool],
        verbose=True,
        backstory=(
            "You are an expert at finding and analyzing AI-related news "
            "from The New York Times. You focus on identifying the most "
            "relevant and recent AI developments, breakthroughs, and industry news."
        )
    )
    
    # Agent 2: Wall Street Journal Scraper
    wsj_agent = Agent(
        role="Wall Street Journal AI Business Analyst",
        goal="Discover AI business news and market trends from The Wall Street Journal",
        tools=[search_tool, scrape_tool],
        verbose=True,
        backstory=(
            "You specialize in AI business news from The Wall Street Journal, "
            "focusing on market trends, investment news, corporate AI adoption, "
            "and the business impact of artificial intelligence technologies."
        )
    )
    
    # Agent 3: Washington Post Scraper
    wapo_agent = Agent(
        role="Washington Post AI Policy Reporter",
        goal="Research AI policy and societal impact news from The Washington Post",
        tools=[search_tool, scrape_tool],
        verbose=True,
        backstory=(
            "You are skilled at finding AI-related news from The Washington Post, "
            "particularly focusing on AI policy, regulation, societal impacts, "
            "and government perspectives on artificial intelligence."
        )
    )
    
    # Agent 4: Blog Writer
    blog_writer = Agent(
        role="AI Blog Writer and Content Synthesizer",
        goal="Create comprehensive blog posts by synthesizing news from multiple sources",
        tools=[],  # No tools needed for writing
        verbose=True,
        backstory=(
            "You are an expert content creator who specializes in synthesizing "
            "news from multiple sources into engaging, informative blog posts. "
            "You excel at identifying common themes, providing analysis, and "
            "creating well-structured content that appeals to AI enthusiasts."
        )
    )
    
    return nyt_agent, wsj_agent, wapo_agent, blog_writer

def create_tasks(nyt_agent, wsj_agent, wapo_agent, blog_writer):
    """Create tasks for each agent"""
    
    # Task 1: NYT News Scraping
    nyt_task = Task(
        description=(
            "Search for the latest AI news from The New York Times. "
            "Look for articles published in the last 7 days about artificial intelligence, "
            "machine learning, ChatGPT, AI companies, or AI breakthroughs. "
            "Analyze and summarize the top 3 most relevant articles."
        ),
        expected_output=(
            "A detailed report with 3 AI news articles from NYT including: "
            "article titles, summaries, key points, URLs, and relevance scores (1-10)."
        ),
        agent=nyt_agent,
        output_json=NewsArticle
    )
    
    # Task 2: WSJ News Scraping
    wsj_task = Task(
        description=(
            "Find the latest AI business news from The Wall Street Journal. "
            "Focus on articles from the past 7 days covering AI investments, "
            "AI company earnings, market trends, and business applications of AI. "
            "Analyze and summarize the top 3 most relevant business articles."
        ),
        expected_output=(
            "A comprehensive report with 3 AI business articles from WSJ including: "
            "article titles, summaries, key business insights, URLs, and relevance scores."
        ),
        agent=wsj_agent,
        output_json=NewsArticle
    )
    
    # Task 3: Washington Post News Scraping
    wapo_task = Task(
        description=(
            "Research the latest AI policy and societal news from The Washington Post. "
            "Look for recent articles about AI regulation, government AI initiatives, "
            "AI ethics, and societal impacts of artificial intelligence. "
            "Analyze and summarize the top 3 most relevant policy articles."
        ),
        expected_output=(
            "A thorough report with 3 AI policy articles from Washington Post including: "
            "article titles, summaries, policy implications, URLs, and relevance scores."
        ),
        agent=wapo_agent,
        output_json=NewsArticle
    )
    
    # Task 4: Blog Writing
    blog_task = Task(
        description=(
            "Using the news articles gathered from NYT, WSJ, and Washington Post, "
            "write a comprehensive blog post about the current state of AI news. "
            "Create an engaging title, introduction, main content with analysis, "
            "and conclusion. Include insights about trends and implications."
        ),
        expected_output=(
            "A well-structured blog post in markdown format with: "
            "engaging title, introduction, main content sections, analysis, "
            "conclusion, and properly cited sources."
        ),
        agent=blog_writer,
        context=[nyt_task, wsj_task, wapo_task],  # Depends on other tasks
        output_file="ai_news_blog.md"
    )
    
    return nyt_task, wsj_task, wapo_task, blog_task

def create_word_document(blog_content, output_path):
    """Convert blog content to Word document"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('AI News Weekly Digest', 0)
    title.alignment = 1  # Center alignment
    
    # Add date
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = 1
    
    # Add content
    lines = blog_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    
    # Save document
    doc.save(output_path)

def create_pdf_document(blog_content, output_path):
    """Convert markdown blog content to PDF"""
    try:
        # Convert markdown to HTML
        html_content = markdown.markdown(blog_content, extensions=['extra', 'codehilite'])
        
        # Add CSS styling for better PDF appearance
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: 'Arial', sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                h1 {{
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    text-align: center;
                }}
                h2 {{
                    color: #34495e;
                    border-bottom: 2px solid #ecf0f1;
                    padding-bottom: 5px;
                    margin-top: 30px;
                }}
                h3 {{
                    color: #7f8c8d;
                    margin-top: 25px;
                }}
                p {{
                    margin-bottom: 15px;
                    text-align: justify;
                }}
                .header {{
                    text-align: center;
                    color: #7f8c8d;
                    font-style: italic;
                    margin-bottom: 30px;
                }}
                ul, ol {{
                    margin-bottom: 15px;
                }}
                li {{
                    margin-bottom: 5px;
                }}
                .footer {{
                    margin-top: 40px;
                    text-align: center;
                    color: #7f8c8d;
                    font-size: 12px;
                    border-top: 1px solid #ecf0f1;
                    padding-top: 20px;
                }}
            </style>
        </head>
        <body>
            <h1>AI News Weekly Digest</h1>
            <div class="header">Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</div>
            {html_content}
            <div class="footer">
                Generated by AI News Blog Generator | Powered by CrewAI Multi-Agent System
            </div>
        </body>
        </html>
        """
        
        # Configure PDF options for better output
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None
        }
        
        # Convert HTML to PDF
        pdfkit.from_string(styled_html, output_path, options=options)
        return True
        
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return False

def create_simple_pdf_fallback(blog_content, output_path):
    """Fallback PDF creation using reportlab if pdfkit fails"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        
        # Create PDF document
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=HexColor('#2c3e50')
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=HexColor('#34495e')
        )
        
        # Build content
        story = []
        
        # Add title
        story.append(Paragraph("AI News Weekly Digest", title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Process content
        lines = blog_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            else:
                story.append(Paragraph(line, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        return True
        
    except ImportError:
        st.error("Neither pdfkit nor reportlab is available for PDF generation.")
        return False
    except Exception as e:
        st.error(f"Error creating PDF with fallback method: {str(e)}")
        return False

def main():
    st.title("🤖 AI News Generator + RAG Query System")
    st.markdown("### Multi-Agent News Scraping + Intelligent Document Querying")
    
    # Sidebar for API keys
    st.sidebar.header("Configuration")
    
    openai_key = st.sidebar.text_input(
        "OpenAI API Key", 
        type="password",
        help="Enter your OpenAI API key"
    )
    
    serper_key = st.sidebar.text_input(
        "Serper API Key", 
        type="password",
        help="Enter your Serper API key for web search"
    )
    
    # Initialize session state
    if 'rag_system' not in st.session_state:
        st.session_state.rag_system = None
    if 'blog_generated' not in st.session_state:
        st.session_state.blog_generated = False
    if 'blog_content' not in st.session_state:
        st.session_state.blog_content = ""
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    
    # Main tabs
    tab1, tab2 = st.tabs(["📰 Generate AI News Blog", "💬 Query Generated Blog"])
    
    with tab1:
        st.subheader("AI News Generation")
        
        if st.button("🚀 Generate AI News Blog", type="primary"):
            if not openai_key or not serper_key:
                st.error("Please provide both API keys!")
                return
            
            # Set environment variables
            os.environ['OPENAI_API_KEY'] = openai_key
            os.environ['SERPER_API_KEY'] = serper_key
            
            # Initialize progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Initialize tools
                status_text.text("Initializing tools...")
                search_tool, scrape_tool = initialize_tools()
                progress_bar.progress(10)
                
                # Create agents
                status_text.text("Creating AI agents...")
                nyt_agent, wsj_agent, wapo_agent, blog_writer = create_agents(search_tool, scrape_tool)
                progress_bar.progress(20)
                
                # Create tasks
                status_text.text("Setting up tasks...")
                nyt_task, wsj_task, wapo_task, blog_task = create_tasks(nyt_agent, wsj_agent, wapo_agent, blog_writer)
                progress_bar.progress(30)
                
                # Create crew
                status_text.text("Assembling the crew...")
                news_crew = Crew(
                    agents=[nyt_agent, wsj_agent, wapo_agent, blog_writer],
                    tasks=[nyt_task, wsj_task, wapo_task, blog_task],
                    verbose=True
                )
                progress_bar.progress(40)
                
                # Execute the crew
                status_text.text("Agents are working... This may take several minutes...")
                
                with st.spinner("🔍 NYT Agent scraping news..."):
                    progress_bar.progress(50)
                
                with st.spinner("💼 WSJ Agent analyzing business news..."):
                    progress_bar.progress(60)
                
                with st.spinner("🏛️ Washington Post Agent researching policy news..."):
                    progress_bar.progress(70)
                
                with st.spinner("✍️ Blog Writer creating content..."):
                    result = news_crew.kickoff()
                    progress_bar.progress(80)
                
                # Debug: Show what CrewAI generated
                st.info(f"🔍 CrewAI working directory: {Path.cwd()}")
                st.info(f"📝 Result type: {type(result)}")
                
                # Create Word and PDF documents
                status_text.text("Creating Word and PDF documents...")
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                word_file_path = output_dir / f"ai_news_blog_{timestamp}.docx"
                pdf_file_path = output_dir / f"ai_news_blog_{timestamp}.pdf"
                
                # Check if CrewAI created the markdown file in current directory
                markdown_files = list(Path.cwd().glob("ai_news_blog.md"))
                if markdown_files:
                    markdown_source = markdown_files[0]
                    st.success(f"✅ Found markdown file: {markdown_source}")
                else:
                    # Check in output directory
                    markdown_source = output_dir / "ai_news_blog.md"
                    if not markdown_source.exists():
                        # Create from result if no file found
                        st.warning("⚠️ No markdown file found, creating from result...")
                        with open(markdown_source, 'w', encoding='utf-8') as f:
                            f.write(str(result))
                
                # Read the markdown content
                with open(markdown_source, 'r', encoding='utf-8') as f:
                    blog_content = f.read()
                
                # Store in session state
                st.session_state.blog_content = blog_content
                st.session_state.blog_generated = True
                
                # Copy markdown to output directory if it's elsewhere
                output_markdown = output_dir / "ai_news_blog.md"
                if markdown_source != output_markdown:
                    with open(output_markdown, 'w', encoding='utf-8') as f:
                        f.write(blog_content)
                
                # Create Word document
                create_word_document(blog_content, word_file_path)
                
                # Create PDF document  
                pdf_created = create_pdf_document(blog_content, pdf_file_path)
                if not pdf_created:
                    # Try fallback method
                    pdf_created = create_simple_pdf_fallback(blog_content, pdf_file_path)
                
                progress_bar.progress(90)
                
                status_text.text("✅ Blog generation completed!")
                progress_bar.progress(100)
                
                # Initialize RAG system with generated content
                if openai_key:
                    client = initialize_openai_client(openai_key)
                    st.session_state.rag_system = RAGSystem(client)
                    if st.session_state.rag_system.build_index(blog_content):
                        st.success("🧠 RAG system initialized! You can now query the generated blog in the Query tab.")
                
                # Display results
                st.success("🎉 AI News Blog Generated Successfully!")
                
                # Debug: List all files in output directory
                files_in_output = list(output_dir.glob("*"))
                st.info(f"📂 Files in output directory ({len(files_in_output)}): {[f.name for f in files_in_output]}")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("📊 Execution Summary")
                    st.write(f"✅ NYT articles analyzed")
                    st.write(f"✅ WSJ business news processed")  
                    st.write(f"✅ Washington Post policy news reviewed")
                    st.write(f"✅ Blog post created")
                    st.write(f"✅ Word document saved: `{word_file_path.name}`")
                    if pdf_file_path.exists():
                        st.write(f"✅ PDF document saved: `{pdf_file_path.name}`")
                
                with col2:
                    st.subheader("📁 Output Files")
                    
                    # PDF Download
                    if pdf_file_path.exists():
                        with open(pdf_file_path, 'rb') as f:
                            st.download_button(
                                label="📄 Download PDF Document",
                                data=f.read(),
                                file_name=pdf_file_path.name,
                                mime="application/pdf"
                            )
                    
                    # Word Download
                    if word_file_path.exists():
                        with open(word_file_path, 'rb') as f:
                            st.download_button(
                                label="📝 Download Word Document",
                                data=f.read(),
                                file_name=word_file_path.name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    
                    # Markdown Download
                    if output_markdown.exists():
                        with open(output_markdown, 'r', encoding='utf-8') as f:
                            markdown_content = f.read()
                        st.download_button(
                            label="📋 Download Markdown File",
                            data=markdown_content,
                            file_name="ai_news_blog.md",
                            mime="text/markdown"
                        )
                
                # Display blog preview
                if output_markdown.exists():
                    st.subheader("📖 Blog Preview")
                    with open(output_markdown, 'r', encoding='utf-8') as f:
                        blog_preview = f.read()
                    st.markdown(blog_preview)
                
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                st.error("Please check your API keys and try again.")
    
    with tab2:
        st.subheader("Query Generated AI News Blog")
        
        if not st.session_state.blog_generated:
            st.warning("⚠️ Please generate a blog first in the 'Generate AI News Blog' tab.")
        else:
            if st.session_state.rag_system is None and openai_key:
                # Initialize RAG system if not already done
                client = initialize_openai_client(openai_key)
                st.session_state.rag_system = RAGSystem(client)
                st.session_state.rag_system.build_index(st.session_state.blog_content)
            
            if st.session_state.rag_system:
                st.success("✅ RAG system ready! Ask questions about the generated AI news blog.")
                
                # Chat interface
                st.subheader("💬 Chat with the AI News Document")
                
                # Display chat history
                for message in st.session_state.chat_messages:
                    with st.chat_message(message["role"]):
                        st.markdown(message["content"])
                
                # Query input
                if query := st.chat_input("Ask a question about the AI news..."):
                    # Add user message to chat history
                    st.session_state.chat_messages.append({"role": "user", "content": query})
                    with st.chat_message("user"):
                        st.markdown(query)
                    
                    # Generate response
                    with st.chat_message("assistant"):
                        with st.spinner("Thinking..."):
                            response = st.session_state.rag_system.generate_response(query)
                        st.markdown(response)
                    
                    # Add assistant response to chat history
                    st.session_state.chat_messages.append({"role": "assistant", "content": response})
                
                # Clear chat button
                if st.button("🗑️ Clear Chat History"):
                    st.session_state.chat_messages = []
                    if st.session_state.rag_system:
                        st.session_state.rag_system.chat_history = []
                    st.rerun()
            else:
                st.error("❌ Please provide OpenAI API key to use the RAG system.")
    
    # Information section
    st.markdown("---")
    st.subheader("ℹ️ How it works")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        **🗞️ NYT Agent**
        - Searches New York Times
        - Finds latest AI news
        - Analyzes tech developments
        """)
    
    with col2:
        st.markdown("""
        **💼 WSJ Agent**
        - Scrapes Wall Street Journal
        - Focuses on AI business news
        - Analyzes market trends
        """)
    
    with col3:
        st.markdown("""
        **🏛️ WaPo Agent**
        - Reviews Washington Post
        - Covers AI policy news
        - Analyzes regulatory trends
        """)
    
    with col4:
        st.markdown("""
        **🧠 RAG System**
        - Indexes generated blog
        - Enables intelligent queries
        - Provides contextual answers
        """)

if __name__ == "__main__":
    main()